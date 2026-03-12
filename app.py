from fastapi import FastAPI, Form, Request
from fastapi.responses import FileResponse, HTMLResponse
from fastapi.staticfiles import StaticFiles
import uvicorn
from docx import Document
from io import BytesIO
import os
import re
import tempfile
from pathlib import Path

app = FastAPI(title="Generador de Acuerdos de Confidencialidad - El Grupo")

# Servir archivos estáticos
app.mount("/static", StaticFiles(directory="static"), name="static")

# Datos fijos de las empresas contratantes
EMPRESAS_CONTRATANTE = {
    "perez_villa": {
        "nombre": "PÉREZ Y VILLA S.A.S",
        "nit": "890.926.395-8",
    },
    "fiera": {
        "nombre": "FIERA S.A.S.",
        "nit": "900.072.392-5",
    }
}

# Datos de contacto fijos del Contratante
CONTACTO_CONTRATANTE = {
    "nombre": "SANTIAGO VÉLEZ GRACIÁN",
    "telefono": "3154719586",
    "direccion": "Carrera 43b # 16 - 95 Manila Of. 614, Medellín, Colombia",
    "correo": "info@elgrupo.com.co"
}


def replace_in_paragraph(paragraph, replacements):
    """
    Reemplaza variables en un párrafo de Word.
    Maneja el caso en que una variable esté repartida en múltiples runs.
    """
    # Intento simple: reemplazar dentro de cada run individual
    for run in paragraph.runs:
        for key, value in replacements.items():
            if key in run.text:
                run.text = run.text.replace(key, value)

    # Verificar si aún quedan variables (pueden estar divididas entre runs)
    full_text = ''.join(run.text for run in paragraph.runs)
    needs_fix = any(key in full_text for key in replacements)

    if needs_fix:
        # Reconstruir el párrafo con el texto reemplazado
        new_text = full_text
        for key, value in replacements.items():
            new_text = new_text.replace(key, value)

        if paragraph.runs:
            # Poner todo el texto en el primer run y limpiar los demás
            first_run = paragraph.runs[0]
            first_run.text = new_text
            for run in paragraph.runs[1:]:
                run.text = ''


def generate_contract(replacements: dict) -> str:
    """Genera el contrato .docx y devuelve la ruta del archivo temporal."""
    template_path = Path("doc_templates/plantilla_acuerdo_confidencialidad.docx")
    doc = Document(str(template_path))

    # Reemplazar en párrafos principales
    for para in doc.paragraphs:
        replace_in_paragraph(para, replacements)

    # Reemplazar en tablas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_in_paragraph(para, replacements)

    # Guardar en archivo temporal
    nombre_contratista = replacements.get("{NOMBRE EMPRESA CONTRATISTA}", "Contratista")
    safe_name = re.sub(r'[^a-zA-Z0-9_\-]', '_', nombre_contratista)
    temp_file = tempfile.NamedTemporaryFile(
        suffix=".docx",
        prefix=f"Acuerdo_Confidencialidad_{safe_name}_",
        delete=False,
        dir="/tmp"
    )
    doc.save(temp_file.name)
    temp_file.close()
    return temp_file.name


@app.get("/", response_class=HTMLResponse)
async def index():
    html_path = Path("static/index.html")
    return HTMLResponse(content=html_path.read_text(encoding="utf-8"))


@app.post("/generar-contrato")
async def generar_contrato(
    empresa_contratante: str = Form(...),
    nombre_empresa_contratista: str = Form(...),
    nit_empresa_contratista: str = Form(...),
    nombre_representante: str = Form(...),
    cedula_representante: str = Form(...),
    direccion_empresa_contratista: str = Form(...),
    telefono_representante: str = Form(...),
    direccion_representante: str = Form(...),
    correo_representante: str = Form(...),
    fecha_firma: str = Form(...),
):
    empresa = EMPRESAS_CONTRATANTE.get(empresa_contratante, EMPRESAS_CONTRATANTE["perez_villa"])

    # Formatear fecha si viene en formato ISO (YYYY-MM-DD)
    try:
        from datetime import datetime
        fecha_obj = datetime.strptime(fecha_firma, "%Y-%m-%d")
        meses = [
            "enero", "febrero", "marzo", "abril", "mayo", "junio",
            "julio", "agosto", "septiembre", "octubre", "noviembre", "diciembre"
        ]
        fecha_formateada = f"{fecha_obj.day} de {meses[fecha_obj.month - 1]} de {fecha_obj.year}"
    except Exception:
        fecha_formateada = fecha_firma

    replacements = {
        "{NOMBRE EMPRESA CONTRATANTE}": empresa["nombre"],
        "{NIT EMPRESA CONTRATANTE}": empresa["nit"],
        "{NOMBRE EMPRESA CONTRATISTA}": nombre_empresa_contratista.upper(),
        "{NIT EMPRESA CONTRATISTA}": nit_empresa_contratista,
        "{NOMBRE REPRESENTANTE LEGAL CONTRATISTA}": nombre_representante.upper(),
        "{CEDULA REPRESTANTE LEGAL EMPRESA CONTRATISTA}": cedula_representante,
        "{DIRECCIÓN EMPRESA CONTRATISTA}": direccion_empresa_contratista,
        "{TELEFONO REPRESENTANTE LEGAL CONTRATISTA}": telefono_representante,
        "{DIRECCIÓN REPRESENTANTE LEGAL CONTRATISTA}": direccion_representante,
        "{CORREO ELECTRÓNICO REPRESENTANTE LEGAL CONTRATISTA}": correo_representante,
        "{FECHA DE FIRMA}": fecha_formateada,
    }

    output_path = generate_contract(replacements)
    filename = f"Acuerdo_Confidencialidad_{nombre_empresa_contratista.replace(' ', '_')}.docx"

    return FileResponse(
        path=output_path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=filename,
        background=None
    )


if __name__ == "__main__":
    port = int(os.environ.get("PORT", 8000))
    uvicorn.run("app:app", host="0.0.0.0", port=port, reload=False)
